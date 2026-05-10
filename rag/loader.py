"""Document loaders: extract plain text from PDF, TXT, MD, DOCX."""
from __future__ import annotations

import io
from pathlib import Path

from pypdf import PdfReader
from docx import Document as DocxDocument


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}


def load_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[Page {i}]\n{text}")
    return "\n\n".join(pages)


def load_txt(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def load_docx(data: bytes) -> str:
    doc = DocxDocument(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def load_document(filename: str, data: bytes) -> str:
    """Dispatch to the right loader based on file extension. Returns plain text."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return load_pdf(data)
    if ext in (".txt", ".md"):
        return load_txt(data)
    if ext == ".docx":
        return load_docx(data)
    raise ValueError(
        f"Unsupported file type: {ext}. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
    )
